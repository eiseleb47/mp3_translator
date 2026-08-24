"""Baut aus Whisper-Segmenten Absätze und schreibt sie als .docx (OnlyOffice-kompatibel)."""

import os
import re
import tempfile
from datetime import datetime
from pathlib import Path

from .audio import format_duration

# lxml lehnt C0-Steuerzeichen ab; Tab und Zeilenumbruch sind erlaubt.
CONTROL_CHARS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def xml_safe(text: str) -> str:
    return CONTROL_CHARS.sub("", text)

PARAGRAPH_GAP = 1.5
SOFT_LIMIT = 400
HARD_LIMIT = 900
SENTENCE_END = (".", "!", "?", "…", ":", '"', "»", "”", ")")

LOCALES = {
    "de": "de-DE", "en": "en-US", "fr": "fr-FR", "it": "it-IT", "es": "es-ES",
    "tr": "tr-TR", "hr": "hr-HR", "sr": "sr-RS", "bs": "bs-BA", "pl": "pl-PL",
    "ro": "ro-RO", "ru": "ru-RU", "ar": "ar-SA", "nl": "nl-NL", "pt": "pt-PT",
}

LANGUAGE_NAMES = {
    "de": "Deutsch", "en": "Englisch", "fr": "Französisch", "it": "Italienisch",
    "es": "Spanisch", "tr": "Türkisch", "hr": "Kroatisch", "sr": "Serbisch",
    "bs": "Bosnisch", "pl": "Polnisch", "ro": "Rumänisch", "ru": "Russisch",
    "ar": "Arabisch", "nl": "Niederländisch", "pt": "Portugiesisch",
}


def group_segments(segments: list[tuple[float, float, str]]) -> list[tuple[float, str]]:
    """Fasst Segmente zu lesbaren Absätzen zusammen (Pausen- und Satzgrenzen)."""
    paragraphs: list[tuple[float, str]] = []
    start: float | None = None
    parts: list[str] = []
    length = 0
    previous_end: float | None = None

    def flush() -> None:
        nonlocal start, parts, length
        if parts:
            paragraphs.append((start or 0.0, " ".join(parts)))
        start, parts, length = None, [], 0

    for seg_start, seg_end, text in segments:
        if parts:
            gap = seg_start - previous_end if previous_end is not None else 0.0
            ends_sentence = parts[-1].endswith(SENTENCE_END)
            if length >= HARD_LIMIT:
                flush()
            elif ends_sentence and (gap >= PARAGRAPH_GAP or length >= SOFT_LIMIT):
                flush()
        if not parts:
            start = seg_start
        parts.append(text)
        length += len(text) + 1
        previous_end = seg_end

    flush()
    return paragraphs


def timestamp(seconds: float) -> str:
    total = int(seconds)
    hours, rest = divmod(total, 3600)
    minutes, secs = divmod(rest, 60)
    if hours:
        return f"[{hours}:{minutes:02d}:{secs:02d}]"
    return f"[{minutes:02d}:{secs:02d}]"


def _set_document_language(document, code: str | None) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if not code:
        return
    locale = LOCALES.get(code, f"{code}-{code.upper()}")

    defaults = document.styles.element.find(qn("w:docDefaults"))
    if defaults is not None:
        rpr_default = defaults.find(qn("w:rPrDefault"))
        if rpr_default is not None:
            rpr = rpr_default.find(qn("w:rPr"))
            if rpr is None:
                rpr = OxmlElement("w:rPr")
                rpr_default.append(rpr)
            lang = rpr.find(qn("w:lang"))
            if lang is None:
                lang = OxmlElement("w:lang")
                rpr.append(lang)
            lang.set(qn("w:val"), locale)

    for style_name in ("Normal", "Title", "Subtitle"):
        try:
            style = document.styles[style_name]
        except KeyError:
            continue
        rpr = style.element.get_or_add_rPr()
        lang = rpr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            rpr.append(lang)
        lang.set(qn("w:val"), locale)


def write_docx(target: Path, title: str, result: dict, source_name: str,
               model: str, with_timestamps: bool = False) -> Path:
    from docx import Document
    from docx.shared import Pt

    document = Document()
    _set_document_language(document, result.get("language"))
    document.styles["Normal"].font.size = Pt(11)
    document.styles["Normal"].paragraph_format.space_after = Pt(8)

    document.add_heading(xml_safe(title), level=1)

    code = result.get("language")
    language_label = LANGUAGE_NAMES.get(code, code or "unbekannt")
    probability = result.get("language_probability")
    if isinstance(probability, (int, float)):
        language_label += f" ({probability * 100:.0f} %)"
    info = (
        f"Quelle: {source_name}  ·  "
        f"Dauer: {format_duration(result.get('duration'))}  ·  "
        f"Sprache: {language_label}  ·  "
        f"Modell: {model}  ·  "
        f"Erstellt: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
    )
    meta = document.add_paragraph()
    run = meta.add_run(xml_safe(info))
    run.italic = True
    run.font.size = Pt(8)

    paragraphs = group_segments(result.get("segments") or [])
    if not paragraphs:
        document.add_paragraph("(Keine Sprache erkannt.)")
    for start, text in paragraphs:
        paragraph = document.add_paragraph()
        if with_timestamps:
            stamp = paragraph.add_run(f"{timestamp(start)} ")
            stamp.bold = True
        paragraph.add_run(xml_safe(text))

    target.parent.mkdir(parents=True, exist_ok=True)
    # Erst vollständig danebenschreiben, dann umbenennen: bricht das Speichern ab
    # (z.B. volle Platte), bleibt keine halbe .docx mit gültigem Namen zurück.
    handle, tmp = tempfile.mkstemp(dir=target.parent, prefix=".docx-", suffix=".tmp")
    os.close(handle)
    try:
        document.save(tmp)
        os.replace(tmp, target)
    except BaseException:
        try:
            os.unlink(tmp)
        except OSError:
            pass
        raise
    return target
